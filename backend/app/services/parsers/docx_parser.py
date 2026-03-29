from pathlib import Path

from docx import Document as DocxDocument

from app.services.parsers.base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    supported_extensions = {".docx"}

    @staticmethod
    def supported_mime_types() -> set[str]:
        return {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }

    def parse(self, path: Path) -> ParsedDocument:
        document = DocxDocument(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        return ParsedDocument(text=text)
